"""
Module pour parser les CVs au format PDF et DOCX.
Includes OCR fallback for scanned documents.
"""

import os
import logging
from typing import Dict, List
from src.parsers.entity_extractor import EntityExtractor

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str, use_ocr: bool = False) -> str:
    """
    Extrait le texte d'un fichier PDF avec fallback OCR.
    
    Args:
        file_path: Chemin vers le fichier PDF
        use_ocr: Si True, utilise OCR si l'extraction standard échoue
    
    Returns:
        str: Texte extrait du PDF
    """
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            text = text.strip()
            
            # Si le texte est trop court et OCR est activé, utiliser OCR
            if use_ocr and len(text) < 100:
                logger.info(f"Texte extrait trop court ({len(text)} chars), utilisation OCR")
                try:
                    from src.parsers.ocr_parser import OCRParser
                    from config.settings import OCRSettings
                    
                    ocr_parser = OCRParser(OCRSettings(enable_ocr=True))
                    result = ocr_parser.process_pdf(file_path)
                    
                    if result['success']:
                        logger.info(f"OCR réussi: {len(result['raw_text'])} chars")
                        return result['raw_text']
                except Exception as e:
                    logger.warning(f"OCR échoué: {e}")
            
            return text
    except Exception as e:
        logger.error(f"Erreur lors de la lecture du PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """
    Extrait le texte d'un fichier DOCX.
    
    Args:
        file_path (str): Chemin vers le fichier DOCX.
    
    Returns:
        str: Texte extrait du DOCX.
    """
    try:
        from docx import Document
        doc = Document(file_path)
        # Extraire le texte des paragraphes et des tableaux
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)
    except Exception as e:
        print(f"Erreur lors de la lecture du DOCX {file_path}: {e}")
        return ""

def parse_cv(cv_path: str) -> Dict[str, str]:
    """
    Parse un CV et extrait le texte brut selon le format.
    OCR is automatically used as fallback for scanned PDFs.
    
    Args:
        cv_path (str): Chemin vers le fichier CV.
    
    Returns:
        Dict[str, str]: Dictionnaire contenant le texte extrait.
    """
    _, ext = os.path.splitext(cv_path)
    text = ""
    if ext.lower() == ".pdf":
        text = extract_text_from_pdf(cv_path, use_ocr=True)
    elif ext.lower() == ".docx":
        text = extract_text_from_docx(cv_path)
    else:
        print(f"Format non supporté : {ext}")
    
    # Extraire les entités structurées
    entity_extractor = EntityExtractor()
    extraction_result = entity_extractor.extract_entities(text)
    
    return {
        "filename": os.path.basename(cv_path),
        "text": text,
        "entities": extraction_result.entities
    }

def load_all_cvs(cv_folder: str) -> List[Dict[str, str]]:
    """
    Charge et parse tous les CVs d'un dossier.
    
    Args:
        cv_folder (str): Chemin vers le dossier contenant les CVs.
    
    Returns:
        List[Dict[str, str]]: Liste des CVs parsés.
    """
    cvs = []
    if not os.path.exists(cv_folder):
        print(f"Le dossier {cv_folder} n'existe pas.")
        return cvs
    
    for file in os.listdir(cv_folder):
        file_path = os.path.join(cv_folder, file)
        if file.lower().endswith(('.pdf', '.docx')):
            cv_data = parse_cv(file_path)
            if cv_data["text"]:
                cvs.append(cv_data)
    return cvs